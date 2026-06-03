import os
import requests
import sys
import json
from docx import Document

def create_test_doc(path):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("甲方联系人：")
    p.add_run("马化腾").bold = True
    p.add_run("，联系电话：13800138000。")
    doc.save(path)

def test_api():
    print("==================================================")
    print("         OpenAPI 自动化集成测试套件")
    print("==================================================")
    
    test_dir = os.path.dirname(os.path.abspath(__file__))
    original_file = os.path.join(test_dir, "openapi_test_original.docx")
    masked_file = os.path.join(test_dir, "openapi_test_masked.docx")
    restored_file = os.path.join(test_dir, "openapi_test_restored.docx")
    key_file = os.path.join(test_dir, "openapi_test_key.json")
    
    # 1. 创建测试原件
    create_test_doc(original_file)
    print(f"[1/5] 创建测试文档成功: {original_file}")
    
    # 2. 调用 OpenAPI 脱敏接口
    print("[2/5] 正在请求 OpenAPI 脱敏接口...")
    url_mask = "http://127.0.0.1:8000/api/openapi/mask"
    
    with open(original_file, "rb") as f:
        files = {"file": (os.path.basename(original_file), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        # 限制只匹配姓名和手机号，验证 categories 过滤
        data = {"categories": "人名,手机号"}
        response = requests.post(url_mask, files=files, data=data)
        
    if response.status_code != 200:
        print(f"[错误] 脱敏接口返回失败，状态码: {response.status_code}, 内容: {response.text}")
        sys.exit(1)
        
    mask_result = response.json()
    print("[成功] 自动化脱敏返回 JSON 响应:")
    print(json.dumps(mask_result, indent=2, ensure_ascii=False))
    
    # 提取映射并导出为备份密钥 JSON
    mappings = mask_result["mappings"]
    with open(key_file, "w", encoding="utf-8") as kf:
        json.dump(mappings, kf, indent=4, ensure_ascii=False)
        
    # 3. 下载脱敏文档
    print("[3/5] 正在通过下载 URL 下载脱敏后的 Word...")
    download_url = f"http://127.0.0.1:8000{mask_result['download_url']}"
    dl_response = requests.get(download_url)
    
    if dl_response.status_code != 200:
        print(f"[错误] 下载脱敏文件失败: {dl_response.status_code}")
        sys.exit(1)
        
    with open(masked_file, "wb") as mf:
        mf.write(dl_response.content)
    print(f"[成功] 脱敏文件下载并保存至: {masked_file}")
    
    # 4. 调用还原接口 (路径 A：自动本地配对数据库)
    print("[4/5] 正在测试路径 A：自动本地配对数据库还原...")
    url_restore = "http://127.0.0.1:8000/api/openapi/restore"
    
    with open(masked_file, "rb") as mf:
        files_restore = {"file": (os.path.basename(masked_file), mf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res_response = requests.post(url_restore, files=files_restore)
        
    if res_response.status_code != 200:
        print(f"[错误] 自动还原失败，状态码: {res_response.status_code}, 内容: {res_response.text}")
        sys.exit(1)
        
    with open(restored_file, "wb") as rf:
        rf.write(res_response.content)
    print(f"[成功] 自动配对还原成功，保存至: {restored_file}")
    
    # 验证还原文档内容
    restored_doc = Document(restored_file)
    restored_text = restored_doc.paragraphs[0].text
    print(f"      还原后的段落内容: {restored_text}")
    if "马化腾" in restored_text and "13800138000" in restored_text:
        print("      内容校验通过！")
    else:
        print("      [错误] 还原文本内容不匹配！")
        sys.exit(1)
        
    # 5. 调用还原接口 (路径 B：外部上传密钥包)
    print("[5/5] 正在测试路径 B：使用外部密钥包还原...")
    # 我们把还原文件删掉，重新生成
    if os.path.exists(restored_file):
        os.remove(restored_file)
        
    with open(masked_file, "rb") as mf, open(key_file, "rb") as kf:
        files_restore_b = {
            "file": (os.path.basename(masked_file), mf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "key_file": (os.path.basename(key_file), kf, "application/json")
        }
        res_response_b = requests.post(url_restore, files=files_restore_b)
        
    if res_response_b.status_code != 200:
        print(f"[错误] 密钥包还原失败，状态码: {res_response_b.status_code}")
        sys.exit(1)
        
    with open(restored_file, "wb") as rf:
        rf.write(res_response_b.content)
    print(f"[成功] 密钥包还原成功，保存至: {restored_file}")
    
    # 再次验证
    restored_doc_b = Document(restored_file)
    restored_text_b = restored_doc_b.paragraphs[0].text
    if "马化腾" in restored_text_b:
        print("      路径 B 校验通过！")
    else:
        print("      [错误] 路径 B 还原内容不匹配！")
        sys.exit(1)
        
    # 清理所有临时测试文件
    for path in [original_file, masked_file, restored_file, key_file]:
        if os.path.exists(path):
            os.remove(path)
            
    print("\n==================================================")
    print("     OpenAPI 所有接口全自动联调测试通过！100% OK")
    print("==================================================")

if __name__ == "__main__":
    test_api()
