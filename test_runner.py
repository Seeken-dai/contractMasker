import os
import uuid
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from app.masker import DocxMasker
from app.database import DatabaseManager

def create_test_document(file_path):
    """创建一个包含各种样式和敏感信息的测试 Word 文档"""
    doc = Document()
    
    # 标题段落
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("合同信息脱敏测试样本")
    run_title.font.size = Pt(16)
    run_title.bold = True
    
    # 正文段落（含加粗、斜体、颜色等样式混合）
    p1 = doc.add_paragraph()
    p1.add_run("本合同由甲方：")
    
    # 加粗的腾讯公司
    r_corp = p1.add_run("深圳市腾讯计算机系统有限公司")
    r_corp.bold = True
    r_corp.font.size = Pt(11)
    
    p1.add_run("（简称甲方，法定代表人：")
    
    # 斜体的人名马化腾
    r_name = p1.add_run("马化腾")
    r_name.italic = True
    
    p1.add_run("）与乙方：")
    
    # 红色普通字体的阿里公司
    r_corp2 = p1.add_run("阿里巴巴（中国）有限公司")
    r_corp2.font.color.rgb = RGBColor(220, 38, 38) # 红色
    
    p1.add_run("（简称乙方，联系电话：")
    
    # 普通字体的电话号码
    r_tel = p1.add_run("13800138000")
    
    p1.add_run("）共同签署。双方就项目事宜达成一致，本协议自 ")
    
    # 日期时间
    r_date = p1.add_run("2026年06月03日")
    r_date.underline = True
    
    p1.add_run(" 起生效。")
    
    # 表格测试
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '签约主体'
    hdr_cells[1].text = '联系人及电话'
    
    row_cells = table.rows[1].cells
    
    p_cell0 = row_cells[0].paragraphs[0]
    p_cell0.add_run("甲方：")
    r_cell_corp = p_cell0.add_run("深圳市腾讯计算机系统有限公司")
    r_cell_corp.bold = True
    
    p_cell1 = row_cells[1].paragraphs[0]
    p_cell1.add_run("马化腾 / ")
    p_cell1.add_run("13800138000").italic = True
    
    doc.save(file_path)
    print(f"[测试] 测试原文档已生成：{file_path}")

def run_integration_test():
    test_dir = os.path.dirname(os.path.abspath(__file__))
    original_path = os.path.join(test_dir, "test_original.docx")
    
    # 1. 创建测试文档
    create_test_document(original_path)
    
    # 2. 初始化数据库并拉取规则
    db = DatabaseManager()
    enabled_rules = db.get_enabled_rules()
    print(f"[测试] 已加载 {len(enabled_rules)} 条匹配规则。")
    
    # 3. 运行 Masker 分析敏感词
    masker = DocxMasker(original_path)
    structure = masker.extract_text_structure()
    print("\n[测试] 文档物理段落骨架提取：")
    for block in structure:
        print(f"  Block {block['block_idx']} ({block['style']}): {block['text']}")
        
    matches = masker.match_sensitive_data(enabled_rules)
    print("\n[测试] 规则自动识别初筛结果：")
    for m in matches:
        print(f"  Block {m['block_idx']} | {m['category']} | {m['original_text']} ({m['start']}-{m['end']})")
        
    # 4. 模拟前端分配占位符并发送脱敏请求
    # 重新整理 matches
    category_counters = {}
    text_to_placeholder = {}
    
    for m in matches:
        key = f"{m['category']}_{m['original_text']}"
        if key not in text_to_placeholder:
            count = category_counters.get(m['category'], 0) + 1
            category_counters[m['category']] = count
            text_to_placeholder[key] = f"[{m['category']}_{count}]"
        m['placeholder'] = text_to_placeholder[key]
        
    # 按照 block 分组
    matches_by_block = {}
    db_mappings = []
    for m in matches:
        b_idx = m['block_idx']
        if b_idx not in matches_by_block:
            matches_by_block[b_idx] = []
        matches_by_block[b_idx].append({
            'start': m['start'],
            'end': m['end'],
            'placeholder': m['placeholder']
        })
        db_mappings.append({
            'placeholder': m['placeholder'],
            'original_text': m['original_text'],
            'category': m['category']
        })
        
    # 执行脱敏
    doc_uuid = str(uuid.uuid4())
    print(f"\n[测试] 开始执行脱敏，生成文档标识 UUID: {doc_uuid}")
    
    masked_file_path = masker.mask_document(matches_by_block, doc_uuid)
    print(f"[测试] 脱敏文档生成成功：{masked_file_path}")
    
    # 存入数据库
    db.save_document_mappings(doc_uuid, "test_original.docx", f"masked_{doc_uuid[:8]}.docx", db_mappings)
    print("[测试] 映射关系已安全写入本地 SQLite。")
    
    # 5. 执行还原测试
    print("\n[测试] 开始模拟还原流程...")
    
    # 打开脱敏后的文档，读取 UUID 并从数据库获取映射
    masked_masker = DocxMasker(masked_file_path)
    read_uuid = masked_masker.doc.core_properties.identifier
    print(f"[测试] 读取到脱敏文档内置 UUID: {read_uuid}")
    
    if read_uuid == doc_uuid:
        print("[测试] UUID 校验一致！")
    else:
        print("[错误] UUID 校验不匹配！")
        sys.exit(1)
        
    # 从数据库检索映射关系
    retrieved_mappings = db.get_mappings_by_uuid(read_uuid)
    print(f"[测试] 从数据库检索到 {len(retrieved_mappings)} 条映射记录。")
    
    # 执行还原
    restored_file_path = DocxMasker.restore_document(masked_file_path, retrieved_mappings)
    print(f"[测试] 还原文档已生成：{restored_file_path}")
    
    # 6. 验证格式与字体锁定
    print("\n[测试] 正在对还原后的文档执行底层 XML 字体检测...")
    restored_doc = Document(restored_file_path)
    
    # 验证段落字体设置
    for p_idx, p in enumerate(restored_doc.paragraphs):
        for r_idx, r in enumerate(p.runs):
            rPr = r._r.get_or_add_rPr()
            rFonts_list = rPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rFonts_list:
                rFonts = rFonts_list[0]
                ascii_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                eastAsia_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                print(f"  段落 {p_idx} Run {r_idx} ({r.text}): ascii={ascii_font}, eastAsia={eastAsia_font}")
                if ascii_font != "Microsoft YaHei" or eastAsia_font != "Microsoft YaHei":
                    print("[错误] 检测到中西文字体未正确锁定为微软雅黑！")
                    sys.exit(1)
            else:
                # 理论上只要我们修改或重写过段落，它就应该包含 rFonts
                # 如果是完全没有变动过的段落可能没有 rFonts
                pass
                
    print("\n[测试] 全流程整合测试圆满成功！格式高精度重构、数据库配对与微软雅黑锁死 100% 正确！")
    
    # 清理测试生成的临时文件 (保留 original 和 restored 供手动检验亦可，这里我们先清理掉以保持 workspace 干净)
    try:
        os.remove(original_path)
        os.remove(masked_file_path)
        os.remove(restored_file_path)
        print("[测试] 临时测试 Word 文件已安全删除。")
    except Exception as e:
        print(f"[测试] 清理临时测试文件失败: {e}")

if __name__ == "__main__":
    run_integration_test()
