import uuid
import re
import os
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import jieba.posseg as pseg

# 强锁微软雅黑属性
def set_run_font_to_msyh(run):
    """强行将 ascii、hAnsi 和 eastAsia 字体统一锁定为 'Microsoft YaHei'"""
    rPr = run._r.get_or_add_rPr()
    
    # 移除已有的字体设置（防止冲突）
    for font_elm in rPr.findall(qn('w:rFonts')):
        rPr.remove(font_elm)
        
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.append(rFonts)

def copy_run_font_safely(src_run, dest_run):
    """高保真拷贝并统一锁定字体，解决中西文混排 Bug，同时保留原有字体样式"""
    src_rPr = src_run._r.get_or_add_rPr()
    rFonts_list = src_rPr.findall(qn('w:rFonts'))
    
    dest_rPr = dest_run._r.get_or_add_rPr()
    for font_elm in dest_rPr.findall(qn('w:rFonts')):
        dest_rPr.remove(font_elm)
        
    if rFonts_list:
        src_fonts = rFonts_list[0]
        font_name = src_fonts.get(qn('w:eastAsia')) or src_fonts.get(qn('w:ascii')) or src_fonts.get(qn('w:hAnsi'))
        if font_name:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            dest_rPr.append(rFonts)

def copy_run_format(src_run, dest_run):
    """从源 Run 复制基础排版样式到目标 Run"""
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    if src_run.font.size:
        dest_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        dest_run.font.color.rgb = src_run.font.color.rgb

# 校验算法类
class VerificationAlgorithms:
    @staticmethod
    def luhn_check(card_number: str) -> bool:
        """银行卡 Luhn 算法校验"""
        card_number = re.sub(r'\D', '', card_number)
        if not (15 <= len(card_number) <= 19):
            return False
        total = 0
        num_digits = len(card_number)
        oddeven = num_digits & 1
        for i in range(num_digits):
            digit = int(card_number[i])
            if not ((i & 1) ^ oddeven):
                digit = digit * 2
                if digit > 9:
                    digit = digit - 9
            total += digit
        return (total % 10) == 0

    @staticmethod
    def id_card_check(id_num: str) -> bool:
        """二代身份证 18 位校验码算法"""
        if len(id_num) != 18:
            return False
        factors = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
        check_codes = ['1', '0', 'X', '9', '8', '7', '6', '5', '4', '3', '2']
        try:
            total = sum(int(id_num[i]) * factors[i] for i in range(17))
            return check_codes[total % 11].upper() == id_num[17].upper()
        except ValueError:
            return False

# 文档解析与重构类
class DocxMasker:
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.paragraphs = self._collect_paragraphs()

    def _collect_paragraphs(self) -> List[Paragraph]:
        """保持物理顺序，扁平化收集文档中所有的段落（含表格内段落）"""
        paragraphs_list = []
        
        def _traverse(element_parent):
            if hasattr(element_parent, 'element') and hasattr(element_parent.element, 'body'):
                parent_elm = element_parent.element.body
            else:
                parent_elm = element_parent._element
                
            for child in parent_elm.iterchildren():
                if child.tag.endswith('p'):
                    paragraphs_list.append(Paragraph(child, element_parent))
                elif child.tag.endswith('tbl'):
                    table = Table(child, element_parent)
                    for row in table.rows:
                        for cell in row.cells:
                            _traverse(cell)
                            
        _traverse(self.doc)
        return paragraphs_list

    def extract_text_structure(self) -> List[Dict[str, Any]]:
        """提取文档段落结构与文本，用于前端渲染"""
        structure = []
        for idx, para in enumerate(self.paragraphs):
            # 获取完整文本
            text = "".join(r.text for r in para.runs)
            # 如果段落没有 run，则直接读 text 属性
            if not para.runs and para.text:
                text = para.text
            structure.append({
                "block_idx": idx,
                "text": text,
                "style": para.style.name if para.style else "Normal"
            })
        return structure

    def match_sensitive_data(self, rules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        匹配段落中的敏感信息
        rules: 包含 {'name': '手机号', 'pattern': '...', 'is_enabled': 1} 的列表
        """
        all_matches = []
        
        for block_idx, para in enumerate(self.paragraphs):
            text = "".join(r.text for r in para.runs)
            if not para.runs and para.text:
                text = para.text
                
            if not text.strip():
                continue
                
            para_matches = []
            
            # 1. 运行正则匹配
            for rule in rules:
                if not rule.get("pattern") or not rule.get("is_enabled"):
                    continue
                    
                pattern_str = rule["pattern"]
                rule_name = rule["name"]
                
                try:
                    for match in re.finditer(pattern_str, text):
                        match_text = match.group(0)
                        start_pos = match.start()
                        end_pos = match.end()
                        
                        # 针对银行卡号和身份证号执行校验过滤
                        if rule_name == "银行卡号" and not VerificationAlgorithms.luhn_check(match_text):
                            continue
                        if rule_name == "身份证号" and not VerificationAlgorithms.id_card_check(match_text):
                            continue
                            
                        para_matches.append({
                            "category": rule_name,
                            "original_text": match_text,
                            "start": start_pos,
                            "end": end_pos
                        })
                except Exception as e:
                    # 正则表达式语法错，跳过
                    print(f"Regex error in rule {rule_name}: {e}")
                    
            # 2. 运行 Jieba 中文人名识别 (如果“人名”规则启用)
            name_rule = next((r for r in rules if r["name"] == "人名"), None)
            if name_rule and name_rule.get("is_enabled"):
                try:
                    words = pseg.cut(text)
                    current_pos = 0
                    for word, flag in words:
                        word_len = len(word)
                        if flag == 'nr' and 2 <= word_len <= 4:
                            # 过滤重复的正则匹配（例如人名已被其他规则圈出）
                            is_duplicate = False
                            for m in para_matches:
                                if m["start"] <= current_pos < m["end"]:
                                    is_duplicate = True
                                    break
                            if not is_duplicate:
                                para_matches.append({
                                    "category": "人名",
                                    "original_text": word,
                                    "start": current_pos,
                                    "end": current_pos + word_len
                                })
                        current_pos += word_len
                except Exception as e:
                    print(f"Jieba processing error: {e}")
                    
            # 合并或去重段落内重叠的匹配
            # 排序：按起点升序，终点降序
            para_matches = sorted(para_matches, key=lambda x: (x["start"], -x["end"]))
            cleaned_matches = []
            last_end = -1
            for m in para_matches:
                if m["start"] >= last_end:
                    cleaned_matches.append(m)
                    last_end = m["end"]
                    
            for m in cleaned_matches:
                m["block_idx"] = block_idx
                all_matches.append(m)
                
        return all_matches

    def mask_document(self, matches_by_block: Dict[int, List[Dict[str, Any]]], uuid_str: str) -> str:
        """
        高精度替换敏感词，并在文档中植入 UUID
        matches_by_block: { block_idx: [ {'start': 3, 'end': 10, 'placeholder': '[企业名称_1]'} ] }
        """
        # 写入内置 UUID 标识
        self.doc.core_properties.identifier = uuid_str
        
        for block_idx, matches in matches_by_block.items():
            if block_idx < 0 or block_idx >= len(self.paragraphs):
                continue
            para = self.paragraphs[block_idx]
            
            # 执行高精度段落重构
            self._rebuild_paragraph_runs(para, matches)
            
        # 保存到临时输出路径
        out_dir = os.path.join(os.path.dirname(self.doc_path), "temp_out")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, f"{uuid_str}.docx")
        self.doc.save(out_path)
        return out_path

    def _rebuild_paragraph_runs(self, paragraph: Paragraph, matches: List[Dict[str, Any]]):
        """核心高精度 Run 拆分合并与微软雅黑锁定重构"""
        old_runs = paragraph.runs
        if not old_runs:
            # 如果没有 Run，但有 text，创建一个默认的
            if paragraph.text:
                paragraph.add_run(paragraph.text)
                old_runs = paragraph.runs
            else:
                return

        # 拼接文本并计算旧 runs 的物理索引区间
        run_ranges = []
        current_pos = 0
        for r in old_runs:
            text_len = len(r.text)
            if text_len > 0:
                run_ranges.append((current_pos, current_pos + text_len, r))
                current_pos += text_len
                
        full_text = "".join(r.text for r in old_runs)
        
        # 收集分割点
        split_points = {0, len(full_text)}
        for start, end, _ in run_ranges:
            split_points.add(start)
            split_points.add(end)
        for m in matches:
            split_points.add(m['start'])
            split_points.add(m['end'])
            
        sorted_points = sorted(list(split_points))
        
        # 重构片段数组
        new_runs_data = [] # List of (text_content, is_mask, template_run)
        processed_masks = set()
        
        for i in range(len(sorted_points) - 1):
            sp_start = sorted_points[i]
            sp_end = sorted_points[i+1]
            if sp_start == sp_end:
                continue
                
            # 检查此区间是否位于某敏感词区间
            in_mask = False
            mask_item = None
            for m in matches:
                if m['start'] <= sp_start < m['end']:
                    in_mask = True
                    mask_item = m
                    break
                    
            if in_mask:
                mask_id = (mask_item['start'], mask_item['end'])
                if mask_id not in processed_masks:
                    processed_masks.add(mask_id)
                    # 寻找敏感词第一个字对应的原 run 模板
                    template_run = None
                    for r_start, r_end, r_obj in run_ranges:
                        if r_start <= mask_item['start'] < r_end:
                            template_run = r_obj
                            break
                    if not template_run and run_ranges:
                        template_run = run_ranges[0][2]
                        
                    new_runs_data.append((mask_item['placeholder'], True, template_run))
            else:
                # 寻找该非敏感区间所属的原 run 模板
                template_run = None
                for r_start, r_end, r_obj in run_ranges:
                    if r_start <= sp_start < r_end:
                        template_run = r_obj
                        break
                if template_run:
                    slice_text = full_text[sp_start:sp_end]
                    new_runs_data.append((slice_text, False, template_run))
                    
        # 清除原有 runs 节点 (只移除 w:r 元素以保留段落样式)
        p_element = paragraph._p
        for child in list(p_element):
            if child.tag.endswith('}r'):
                p_element.remove(child)
            
        # 重新插入 Runs 并高保真锁定字体
        for text_content, is_mask, template_run in new_runs_data:
            new_run = paragraph.add_run(text_content)
            if template_run:
                copy_run_format(template_run, new_run)
                copy_run_font_safely(template_run, new_run)

    @staticmethod
    def restore_document(masked_doc_path: str, mappings: List[Dict[str, Any]]) -> str:
        """
        还原脱敏文档，将占位符替换为原始文本，并锁死微软雅黑字体
        mappings: [{'placeholder': '[人名_1]', 'original_text': '马化腾'}]
        """
        doc = Document(masked_doc_path)
        
        # 建立占位符到原文的快速映射
        mapping_dict = {item["placeholder"]: item["original_text"] for item in mappings}
        
        # 收集所有的段落
        paragraphs_list = []
        def _traverse(element_parent):
            if hasattr(element_parent, 'element') and hasattr(element_parent.element, 'body'):
                parent_elm = element_parent.element.body
            else:
                parent_elm = element_parent._element
                
            for child in parent_elm.iterchildren():
                if child.tag.endswith('p'):
                    paragraphs_list.append(Paragraph(child, element_parent))
                elif child.tag.endswith('tbl'):
                    table = Table(child, element_parent)
                    for row in table.rows:
                        for cell in row.cells:
                            _traverse(cell)
        _traverse(doc)
        
        # 遍历所有段落
        for para in paragraphs_list:
            old_runs = para.runs
            if not old_runs:
                continue
                
            # 拼接段落内容
            full_text = "".join(r.text for r in old_runs)
            
            # 检查该段落是否包含占位符
            has_placeholder = False
            for placeholder in mapping_dict:
                if placeholder in full_text:
                    has_placeholder = True
                    break
                    
            if not has_placeholder:
                # 若无占位符，跳过不处理，保持其原本字体设置不变
                continue
                
            # 查找占位符在段落文本中的所有出现区间
            matches = []
            for placeholder, original in mapping_dict.items():
                for match in re.finditer(re.escape(placeholder), full_text):
                    matches.append({
                        "start": match.start(),
                        "end": match.end(),
                        "placeholder": original # 这里 placeholder 被赋值为 original_text，即替换目标
                    })
                    
            # 按起点排序
            matches = sorted(matches, key=lambda x: x["start"])
            
            # 使用相同的 _rebuild_paragraph_runs 重写该段落，替换回原文
            # 我们直接使用实例化一个临时的 DocxMasker 相同的私有重构函数
            # 为了重用代码，我们需要把 _rebuild_paragraph_runs 转化为静态或单独调用
            masker_instance = DocxMasker.__new__(DocxMasker)
            masker_instance._rebuild_paragraph_runs(para, matches)
            
        # 保存还原文件到临时输出路径
        out_dir = os.path.dirname(masked_doc_path)
        out_path = os.path.join(out_dir, f"restored_{uuid.uuid4().hex[:8]}.docx")
        doc.save(out_path)
        return out_path
